from flask import Flask, render_template, request, redirect, flash, url_for
from werkzeug.utils import secure_filename
import os
from docx import Document
import Create_invoice

app = Flask(__name__)
app.secret_key = 'supersecretkey'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['TEMPLATES_FOLDER'] = 'templates/invoices/'

# Ensure the templates folder exists
if not os.path.exists(app.config['TEMPLATES_FOLDER']):
    os.makedirs(app.config['TEMPLATES_FOLDER'])

# Load the blank template for invoices
def load_invoice_template():
    template_path = '/Users/hsenhsnen/Desktop/Leather/templates/invoices/invoice_template.docx'
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Invoice template not found at {template_path}.")  # Debug line
    return Document(template_path)





@app.route('/create_invoice', methods=['GET', 'POST'])
def create_invoice():
    if request.method == 'POST':
        try:
            # Fill in the invoice data
            invoice_data = {
                'invoice_number': request.form['invoice_number'],
                'date': request.form['date'],
                'buyer': request.form['buyer'],
                'products': request.form.getlist('product'),
                'quantities': list(map(int, request.form.getlist('quantity'))),  # Convert to integers
                'prices': list(map(float, request.form.getlist('price'))),  # Convert to floats
                # Add more fields as necessary
            }

            # Load the blank template and modify it
            document = load_invoice_template()
            for para in document.paragraphs:
                if "PERFORMA INVOICE" in para.text:
                    para.text = f"PERFORMA INVOICE {invoice_data['invoice_number']} /24"
                if "Date" in para.text:
                    para.text = f"{invoice_data['date']}"
                if "buyer:" in para.text:
                    para.text = f"Buyer: {invoice_data['buyer']}"
            
            # Adding product rows to the table
            table = document.tables[0]  # Assuming the first table is for products
            for i in range(len(invoice_data['products'])):
                row_cells = table.add_row().cells
                row_cells[0].text = invoice_data['products'][i]
                row_cells[1].text = str(invoice_data['quantities'][i])
                row_cells[2].text = str(invoice_data['quantities'][i])  # Assuming average weight = quantity
                row_cells[3].text = str(invoice_data['quantities'][i])  # Assuming gross weight = quantity
                row_cells[4].text = "0"  # Placeholder for salt discount
                row_cells[5].text = str(invoice_data['quantities'][i])  # Assuming total weight = quantity
                row_cells[6].text = f"{invoice_data['prices'][i]:.2f}"
                row_cells[7].text = f"{(invoice_data['quantities'][i] * invoice_data['prices'][i]):.2f}"

            # Calculate total amounts for each column (you might want to sum these up)
            total_quantity = sum(invoice_data['quantities'])
            total_price = sum(invoice_data['quantities'][i] * invoice_data['prices'][i] for i in range(len(invoice_data['products'])))
            
            # Add total row
            total_row = document.add_paragraph('Total:')
            total_row.add_run(f" {total_quantity}").bold = True  # Total Quantity
            total_row.add_run(f" {total_price:.2f}").bold = True  # Total Amount

            # Save the new invoice
            saved_invoice_filename = f"invoice_{invoice_data['invoice_number']}.docx"
            invoice_path = os.path.join(app.config['UPLOAD_FOLDER'], saved_invoice_filename)
            document.save(invoice_path)
            flash(f'Invoice {invoice_data["invoice_number"]} created and saved as {saved_invoice_filename}')
            return redirect(url_for('manage_invoices'))
        except Exception as e:
            flash(f'An error occurred: {str(e)}')  # Show error message
            return redirect(url_for('create_invoice'))

    return render_template('create_invoice.html')


# Route for managing invoices
@app.route('/invoices', methods=['GET'])
def manage_invoices():
    # List all generated invoices
    uploaded_invoices = os.listdir(app.config['UPLOAD_FOLDER'])
    return render_template('invoice.html', invoices=uploaded_invoices)

# Default route to redirect to dashboard
@app.route('/')
def dashboard():
    return render_template('dashboard.html')

if __name__ == '__main__':
    app.run(debug=True)
