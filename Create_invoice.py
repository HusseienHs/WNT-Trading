from docx import Document
from docx.shared import Inches


def create():
    # Create a new Document
    doc = Document()

    # Title
    doc.add_heading('PERFORMA INVOICE', level=1)
    doc.add_paragraph('*** / 24')
    doc.add_paragraph('**.**.24')
    doc.add_paragraph('Buyer:')
    doc.add_paragraph('************')

    # Table for Products
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'

    # Adding header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Quantity pcs'
    hdr_cells[2].text = 'Average Weight kg / pcs'
    hdr_cells[3].text = 'Gross weight kg'
    hdr_cells[4].text = 'Salt discount 3%'
    hdr_cells[5].text = 'Total weigh net kg'
    hdr_cells[6].text = 'Price $ / KG'
    hdr_cells[7].text = 'Total Amount $'

    # Example of adding product rows (replace *** with actual values)
    for _ in range(2):  # Adjust number of products as needed
        row_cells = table.add_row().cells
        row_cells[0].text = 'Raw salted hides(***)'  # Product name
        row_cells[1].text = '***'  # Quantity pcs
        row_cells[2].text = '***'  # Average Weight kg / pcs
        row_cells[3].text = '***'  # Gross weight kg
        row_cells[4].text = '***'  # Salt discount 3%
        row_cells[5].text = '***'  # Total weigh net kg
        row_cells[6].text = '***'  # Price $ / KG
        row_cells[7].text = '***'  # Total Amount $

    # Total row
    total_row = doc.add_paragraph('Total:')
    total_row.add_run(' ***').bold = True  # Total Quantity
    total_row.add_run(' ***').bold = True  # Total Average Weight
    total_row.add_run(' ***').bold = True  # Total Gross weight
    total_row.add_run(' ***').bold = True  # Total Salt discount
    total_row.add_run(' ***').bold = True  # Total net weight
    total_row.add_run(' ***').bold = True  # Total price per KG
    total_row.add_run(' ***').bold = True  # Total amount

    # Declaration
    doc.add_heading('Declaration:', level=2)
    declarations = [
        '1. Hides are origin of Israel slaughter house',
        '2. Hides are machine cut.',
        '3. Hides are 80% European breed & 20% Australian',
        '4. Port of loading: Haifa port',
        '5. Payment terms: Payment by LC 45 days',
        '6. Total amount of shipment: ****$'
    ]
    for declaration in declarations:
        doc.add_paragraph(declaration)

    # Bank Account Information
    doc.add_heading('Our Bank Account:', level=2)
    bank_info = [
        'Company name: W.N.T TRADING OR LTD',
        'Company address: I’bilin – ISRAEL pob: 0508',
        'Bank name: MERCANTILE DISCOUNT BANK LTD',
        'Bank address: IBILIN BRANCH 0756, MAIN STR.',
        'Account number: 9073614',
        'Swift: BARDILIT',
        'Iban: IL620177560000090736143'
    ]
    for info in bank_info:
        doc.add_paragraph(info)

    # Save the document
    doc.save('Proforma_Invoice.docx')

create()
